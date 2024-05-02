# Modules for dpp
import numpy as np
import pandas as pd
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split

# Modules for nlp
import nltk
import string
import re
import pandas as pd
import docx
import json
import unicodedata
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
from gensim.models import Word2Vec

nltk.download('punk')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')


# Data_Preprocessing
class data_preprocessing:

    def poz_read_file(uploaded_file):
        # Check if the uploaded file is None
        if uploaded_file is None:
            return None

        # Read the file content
        file_extension = uploaded_file.name.split('.')[-1].lower()
        if file_extension == 'csv':
            return pd.read_csv(uploaded_file)
        elif file_extension in ['xls', 'xlsx']:
            return pd.read_excel(uploaded_file)
        elif file_extension == 'txt':
            data = uploaded_file.getvalue().decode("utf-8")
            return pd.DataFrame({'text_data': [data]})
        else:
            print("Unsupported file type.")
            return None

    # def poz_handle_missing_values(df, fill_value=None):
    #     """
    #     Handle missing values in a DataFrame using a default method.
    #
    #
    #    Parameters:
    #     df (pandas.DataFrame): The DataFrame containing missing values.
    #     fill_value (int, float, str): The value used for filling missing values when method is 'constant'.
    #
    #
    #    Returns:
    #     pandas.DataFrame: The DataFrame with missing values handled using the default method.
    #     """
    #     try:
    #         # If fill_value is not provided, use mean as the default method
    #         if fill_value is None:
    #             fill_value = df.mean()
    #
    #         df_filled = df.fillna(fill_value)
    #     except Exception as e:
    #         print(f"An error occurred: {e}")
    #         df_filled = pd.DataFrame()  # Return an empty DataFrame in case of error
    #
    #     return df_filled

    def poz_handle_missing_values(df, method='mean'):
        """
        Handle missing values in a DataFrame using a specified method.

        Parameters:
        df (pandas.DataFrame): The DataFrame containing missing values.
        method (str): The method used for filling missing values. Options: 'mean', 'median', 'mode', 'constant'.

        Returns:
        pandas.DataFrame: The DataFrame with missing values handled using the specified method.
        """
        try:
            if method == 'mean':
                fill_value = df.mean()
            elif method == 'median':
                fill_value = df.median()
            elif method == 'mode':
                fill_value = df.mode().iloc[0]  # Get the mode value
            elif method == 'constant':
                # You can specify a constant value to fill missing values
                fill_value = 0  # Change this to the desired constant value
            else:
                raise ValueError("Invalid method. Supported methods: 'mean', 'median', 'mode', 'constant'")

            df_filled = df.fillna(fill_value)
        except Exception as e:
            print(f"An error occurred: {e}")
            df_filled = pd.DataFrame()  # Return an empty DataFrame in case of error

        return df_filled

    def poz_standardize(df):
        """
        Standardize numerical data in a DataFrame using z-score normalization.

        Parameters:
        df (pandas.DataFrame): The DataFrame containing numerical data.

        Returns:
        pandas.DataFrame: The DataFrame with numerical data standardized.
        """
        try:
            # Select numerical columns
            numerical_columns = df.select_dtypes(include='number')

            # Standardize numerical columns using z-score normalization
            standardized_df = (numerical_columns - numerical_columns.mean()) / numerical_columns.std()

            # Combine standardized numerical columns with non-numeric columns
            for column in df.columns:
                if column in numerical_columns.columns:
                    df[column] = standardized_df[column]
        except Exception as e:
            print(f"An error occurred: {e}")
            df = pd.DataFrame()  # Return an empty DataFrame in case of error

        return df

    def poz_normalize(df):
        """
        Normalize numerical data in a DataFrame using min-max normalization.

        Parameters:
        df (pandas.DataFrame): The DataFrame containing numerical data.

        Returns:
        pandas.DataFrame: The DataFrame with numerical data normalized.
        """
        try:
            # Select numerical columns
            numerical_columns = df.select_dtypes(include='number')

            # Normalize numerical columns using min-max normalization
            normalized_df = (numerical_columns - numerical_columns.min()) / (
                    numerical_columns.max() - numerical_columns.min())

            # Combine normalized numerical columns with non-numeric columns
            for column in df.columns:
                if column in numerical_columns.columns:
                    df[column] = normalized_df[column]
        except Exception as e:
            print(f"An error occurred: {e}")
            df = pd.DataFrame()  # Return an empty DataFrame in case of error

        return df

    def poz_apply_pca(data, n_components):
        try:
            # Instantiate PCA with desired number of components
            pca = PCA(n_components=n_components)
            # Fit PCA model to data and transform data to lower-dimensional space
            transformed_data = pca.fit_transform(data)
            return transformed_data
        except Exception as e:
            print(f"An error occurred during PCA: {e}")
            return None

    def poz_data_integration(datasets):
        try:
            integrated_data = np.concatenate(datasets, axis=1)
            return integrated_data
        except Exception as e:
            print(f"An error occurred during data integration: {e}")
            return None

    def poz_data_aggregation(data, group_by_columns, aggregation_functions):
        """
        Aggregate data based on specified grouping columns and aggregation functions.

        Parameters:
        - data: DataFrame containing the dataset.
        - group_by_columns: List of columns to group the data by.
        - aggregation_functions: Dictionary where keys are column names and values are aggregation functions.

        Returns:
        - DataFrame with aggregated data.
        """
        try:
            aggregated_data = data.groupby(group_by_columns).agg(aggregation_functions).reset_index()
            return aggregated_data
        except Exception as e:
            print(f"An error occurred during data aggregation: {e}")
            return None

    def poz_split_data(data, target_column, test_size, random_state, axis):
        """
        Split data into train and test sets.


       Parameters:
        - data: DataFrame containing the dataset.
        - target_column: Name of the target column.
        - test_size: Size of the test set.
        - random_state: Random state for reproducibility.
        - axis: Axis along which to drop the target column (0 for rows, 1 for columns).


       Returns:
        - X_train, X_test, y_train, y_test: Train and test sets for features and target.
        """
        if axis not in [0, 1]:
            raise ValueError("Axis must be 0 or 1.")

        X = data.drop(target_column, axis=axis)
        y = data[target_column]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=random_state)
        # data = [X_train, X_test, y_train, y_test]
        return [X_train, X_test, y_train, y_test]

    def poz_outliers(data, method='z-score', threshold=3, action='remove'):
        """
        Detect and handle outliers in a dataset.

        Parameters:
        - data (DataFrame): The input DataFrame containing the data to be processed.
        - method (str): The method for detecting outliers. Options: 'z-score', 'iqr'. Default is 'z-score'.
        - threshold (float): The threshold value for outlier detection. Default is 3.
        - action (str): The action to take for handling outliers. Options: 'remove', 'transform', 'cap'. Default is 'remove'.

        Returns:
        - cleaned_data (DataFrame): The DataFrame with outliers handled based on the specified action.
        """
        cleaned_data = data.copy()

        if method == 'z-score':
            # Detect outliers using z-score
            z_scores = (cleaned_data - cleaned_data.mean()) / cleaned_data.std()
            outliers = abs(z_scores) > threshold

        elif method == 'iqr':
            # Detect outliers using interquartile range (IQR)
            q1 = cleaned_data.quantile(0.25)
            q3 = cleaned_data.quantile(0.75)
            iqr = q3 - q1
            lower_bound = q1 - threshold * iqr
            upper_bound = q3 + threshold * iqr
            outliers = (cleaned_data < lower_bound) | (cleaned_data > upper_bound)

        else:
            raise ValueError("Invalid method. Options: 'z-score', 'iqr'")

        if action == 'remove':
            # Remove outliers
            cleaned_data = cleaned_data[~outliers.any(axis=1)]

        elif action == 'transform':
            # Transform outliers to the mean value
            cleaned_data[outliers] = cleaned_data.mean()

        elif action == 'cap':
            # Cap outliers to a specified range
            cleaned_data[outliers] = cleaned_data.clip(lower=lower_bound, upper=upper_bound, axis=1)

        else:
            raise ValueError("Invalid action. Options: 'remove', 'transform', 'cap'")

        return cleaned_data

    @classmethod
    def poz_remove_char(cls, uploaded_file):
        pass


# Data_Preprocessing for NLP
class NLP:
    def poz_tokenize_text(text):
        """
        Tokenize a text string into words using NLTK word tokenizer.

        Parameters:
        text (str): The input text string to be tokenized.

        Returns:
        list: A list of tokens (words) extracted from the input text.
        """
        import nltk
        nltk.download('punkt')
        tokens = word_tokenize(text)

        return tokens

    def poz_lowercase_text(text):
        # Convert text to lowercase using the lower() method
        lowercased_text = text.lower()

        return lowercased_text

    def poz_remove_punctuation(text):
        """
        Remove punctuation from a text string.

        Parameters:
        text (str): The input text string.

        Returns:
        str: The input text string with punctuation removed.
        """
        # Define a translation table that maps all punctuation characters to None
        translation_table = str.maketrans('', '', string.punctuation)

        # Remove punctuation using the translation table
        text_without_punctuation = text.translate(translation_table)

        return text_without_punctuation

    def poz_stemming_text(text):
        """
        Perform stemming on a text string.

        Parameters:
        text (str): The input text string.

        Returns:
        str: The input text string with words stemmed.
        """
        # Download NLTK resources if not already downloaded
        nltk.download('punkt', quiet=True)

        # Initialize Porter Stemmer
        stemmer = PorterStemmer()

        # Tokenize the input text string
        words = word_tokenize(text)

        # Stem each word in the tokenized words
        stemmed_words = [stemmer.stem(word) for word in words]

        # Join the stemmed words back into a string
        stemmed_text = ' '.join(stemmed_words)

        return stemmed_text

    def poz_expand_contractions(text, contraction_mapping):
        """
        Expand contractions in a text using a predefined contraction mapping.

        Parameters:
        text (str): The input text string.
        contraction_mapping (dict): A dictionary mapping contractions to their expanded forms.

        Returns:
        str: The input text string with contractions expanded.
        """
        # Iterate over each contraction in the mapping and replace it with its expanded form
        for contraction, expanded_form in contraction_mapping.items():
            text = text.replace(contraction, expanded_form)

        return text

    def poz_remove_char(file_path):
        if file_path.endswith('.csv'):
            # Read CSV file
            df = pd.read_csv(file_path)
            # Concatenate all text columns (assuming they're all strings)
            text = ' '.join(df.select_dtypes(include=['object']).stack().tolist())
        elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
            # Read Excel file
            df = pd.read_excel(file_path)
            # Concatenate all text columns (assuming they're all strings)
            text = ' '.join(df.select_dtypes(include=['object']).stack().tolist())
        elif file_path.endswith('.txt'):
            # Read TXT file
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif file_path.endswith('.doc') or file_path.endswith('.docx'):
            # Read DOC or DOCX file
            doc = docx.Document(file_path)
            # Extract text from paragraphs
            text = ' '.join([para.text for para in doc.paragraphs])
        elif file_path.endswith('.json'):
            # Read JSON file
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            # Concatenate all values (assuming they're all strings)
            text = ' '.join(str(value) for value in json_data.values())
        else:
            raise ValueError("Unsupported file format")

        # Remove numerical digits and special characters
        text = re.sub(r'\d+', '', text)  # Remove digits
        text = re.sub(r'[^\w\s]', '', text)  # Remove special characters

        return text

    def remove_noisy_text(text):
        # Remove URLs
        text = re.sub(r'http\S+|www\S+', '', text)

        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)

        # Remove other noisy text patterns as needed

        return text

    def poz_remove_noisy_text(file_path):
        if file_path.endswith('.csv'):
            # Read CSV file
            df = pd.read_csv(file_path)
            # Concatenate all text columns (assuming they're all strings)
            text = ' '.join(df.select_dtypes(include=['object']).stack().tolist())
        elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
            # Read Excel file
            df = pd.read_excel(file_path)
            # Concatenate all text columns (assuming they're all strings)
            text = ' '.join(df.select_dtypes(include=['object']).stack().tolist())
        elif file_path.endswith('.txt'):
            # Read TXT file
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif file_path.endswith('.docx') or file_path.endswith('.doc'):
            # Read DOCX or DOC file
            doc = docx.Document(file_path)
            # Extract text from paragraphs
            text = ' '.join([para.text for para in doc.paragraphs])
        elif file_path.endswith('.json'):
            # Read JSON file
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            # Concatenate all values (assuming they're all strings)
            text = ' '.join(str(value) for value in json_data.values())
        else:
            raise ValueError("Unsupported file format")

        # Remove noisy text patterns
        text = NLP.remove_noisy_text()

        return text

    def poz_word_embed(text):
        # Tokenize the text into words
        words = word_tokenize(text)

        # Train Word2Vec model
        model = Word2Vec([words], vector_size=100, window=5, min_count=1, workers=4).wv['example']

        return model

    def poz_padding(sequence, max_length):
        """
        Pad sequences to a fixed length with zeros.

        Args:
        - sequence (list): Input sequence.
        - max_length (int): Maximum length of the padded sequence.

        Returns:
        - padded_sequence (list): Padded sequence.
        """
        padded_sequence = sequence[:max_length] + [0] * (max_length - len(sequence))
        return padded_sequence

    def poz_truncation(sequence, max_length):
        """
        Truncate sequences to a fixed length.

        Args:
        - sequence (list): Input sequence.
        - max_length (int): Maximum length of the truncated sequence.

        Returns:
        - truncated_sequence (list): Truncated sequence.
        """
        truncated_sequence = sequence[:max_length]
        return truncated_sequence

    def poz_nlp_normalize(text):
        """
        Normalize text by removing accents.

        Args:
        - text (str): Input text to be normalized.

        Returns:
        - normalized_text (str): Normalized text without accents.
        """
        # Remove accents
        normalized_text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8', 'ignore')

        return normalized_text

    def poz_name_entity_recognition(text):
        """
        Perform Named Entity Recognition (NER) on the input text using NLTK.

        Args:
        - text (str): Input text.

        Returns:
        - entities (list): List of named entities detected in the text.
        """
        # Tokenize the text
        tokens = nltk.word_tokenize(text)

        # Perform part-of-speech tagging
        pos_tags = nltk.pos_tag(tokens)

        # Perform named entity recognition
        entities = nltk.ne_chunk(pos_tags)

        return entities

    def poz_part_of_speech(text):
        """
        Perform Part-of-Speech (POS) tagging on the input text using NLTK.

        Args:
        - text (str): Input text.

        Returns:
        - pos_tags (list): List of tuples containing words and their corresponding POS tags.
        """
        # Tokenize the text
        tokens = nltk.word_tokenize(text)

        # Perform part-of-speech tagging
        pos_tags = nltk.pos_tag(tokens)

        return pos_tags
